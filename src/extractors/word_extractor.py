"""
word_extractor.py

WordExtractor — implements docs/06_word_extractor_design.md.

Extracts operational data from `Parte de Operaciones.docx` and
produces a single structured dataset (OperationsDataset), one record
per route segment.

Forward-fill note (see docs/06_word_extractor_design.md, "Data
Normalization Considerations"): earlier drafts of the design assumed
the Train Number cell would be blank on non-first rows of a segment
group, requiring a forward-fill step. That assumption was checked
against the real .docx and does not hold — the Train Number cells are
vertically merged at the document level, so `python-docx` already
returns the correct value on every row (verified: same underlying XML
cell object across merged rows). No forward-fill is implemented here
or anywhere downstream, because there is nothing to fill.

Extraction only — no validation, no comparison against the PDF dataset,
no Excel writes (all explicitly out of scope per the design doc).
"""

from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document


@dataclass
class OperationsRecord:
    """One row of the OperationsDataset — one route segment."""

    route_segment: str
    train_number: str
    tickets_sold: int
    occupancy_pct: str  # kept as-is (e.g. "14.9%") — not required by the design doc,
    # but free to extract since it's already on the row; not used for
    # writing to Programa.xlsx, only potentially useful for validation
    # range checks later.


@dataclass
class WordExtractionResult:
    operations: list[OperationsRecord] = field(default_factory=list)
    skipped_rows: list[str] = field(default_factory=list)


EXPECTED_HEADER = ["Ruta", "Número de tren", "Cantidad de boletos vendidos", "% de Ocupación"]


class WordExtractor:
    """
    Implements docs/06_word_extractor_design.md.

    Extraction only — no validation, no business rules, no Excel
    writes (all explicitly out of scope per the design doc).
    """

    def extract(self, docx_path: str) -> WordExtractionResult:
        result = WordExtractionResult()
        doc = Document(docx_path)

        if not doc.tables:
            raise ValueError(f"No tables found in {docx_path}")

        table = doc.tables[0]
        rows = table.rows

        # First row is the header — skip it, but keep it for traceability
        # in case the document's column order ever changes unexpectedly.
        header = [c.text.strip() for c in rows[0].cells]
        if header != EXPECTED_HEADER:
            result.skipped_rows.append(f"HEADER MISMATCH: {header}")

        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                result.skipped_rows.append(" | ".join(cells))
                continue

            route_segment, train_number, tickets_sold_raw, occupancy_pct = cells[:4]

            if not route_segment or not train_number:
                result.skipped_rows.append(" | ".join(cells))
                continue

            try:
                tickets_sold = int(tickets_sold_raw)
            except ValueError:
                result.skipped_rows.append(" | ".join(cells))
                continue

            result.operations.append(
                OperationsRecord(
                    route_segment=route_segment,
                    train_number=train_number,
                    tickets_sold=tickets_sold,
                    occupancy_pct=occupancy_pct,
                )
            )

        return result


if __name__ == "__main__":
    import sys

    path = sys.argv[1] if len(sys.argv) > 1 else "data/raw/parte_operaciones_sample.docx"
    extractor = WordExtractor()
    res = extractor.extract(path)

    print(f"OperationsDataset: {len(res.operations)} records")
    for rec in res.operations:
        print(
            f"  train={rec.train_number:5s} tickets={rec.tickets_sold:4d}  "
            f"occupancy={rec.occupancy_pct:7s} route={rec.route_segment}"
        )

    print(f"\nSkipped rows: {len(res.skipped_rows)}")
    for line in res.skipped_rows:
        print(f"  {line}")
